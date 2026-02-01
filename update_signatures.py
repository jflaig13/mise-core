#!/usr/bin/env python3
"""
Script to update organizer signature sections in entertainment agreement documents.
"""

from docx import Document
from docx.shared import Inches
import os

# Define the base directory and files
base_dir = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026"
signature_image = "/Users/jonathanflaig/Signature/Final_JonFlaig_Signature.png"

files = [
    "Entertainment_Agreement_Alex_Napier_Band_2026.docx",
    "Entertainment_Agreement_Black_Creek_String_Band_2026_1.docx",
    "Entertainment_Agreement_Blues_Old_Stand_Trio_2026.docx",
    "Entertainment_Agreement_Boukou_Groove_2026.docx",
    "Entertainment_Agreement_Dawghaus_Jazz_Band_2026.docx",
    "Entertainment_Agreement_Eric_Stockton_Band_2026.docx",
    "Entertainment_Agreement_Goodbye_To_Sunshine_2026.docx",
    "Entertainment_Agreement_John_Hart_Project_2026.docx",
    "Entertainment_Agreement_Killer_Robot_Army_2026_1.docx",
    "Entertainment_Agreement_Lips_Manly_2026.docx",
    "Entertainment_Agreement_Shenaniganza_2026.docx",
    "Entertainment_Agreement_Stevie_Monce_2026.docx",
    "Entertainment_Agreement_The_Unnamed_Band_2026.docx",
]

def find_organizer_signature_paragraph(doc):
    """
    Find the paragraph that contains the organizer signature.
    It's after "On behalf of Bloom's Restaurant, LLC d/b/a Papa Surf Burger Bar:"
    and before "Name: Jon Flaig"
    """
    # First, find "Name: Jon Flaig" paragraph
    jon_flaig_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "Name: Jon Flaig" in para.text:
            jon_flaig_idx = i
            break

    if jon_flaig_idx is None:
        return None

    # Now find the organizer line and look for the signature paragraph between them
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if "On behalf of Bloom's Restaurant, LLC d/b/a Papa Surf Burger Bar:" in text:
            # The signature paragraph is the one right before "Name: Jon Flaig"
            # It should contain "Date:" or be between this line and Jon Flaig
            if i < jon_flaig_idx:
                # Look for the paragraph with "Date:" between organizer line and Jon Flaig
                for j in range(i + 1, jon_flaig_idx):
                    if "Date:" in doc.paragraphs[j].text or j == jon_flaig_idx - 1:
                        return doc.paragraphs[j]

    return None

def update_signature(file_path, signature_img_path):
    """
    Update the organizer signature section in a document.
    """
    print(f"\nProcessing: {os.path.basename(file_path)}")

    # Load the document
    doc = Document(file_path)

    # Find the signature paragraph
    sig_para = find_organizer_signature_paragraph(doc)

    if sig_para is None:
        print(f"  ERROR: Could not find organizer signature paragraph")
        return False

    print(f"  Found signature paragraph: '{sig_para.text[:50]}...'")

    # Clear the paragraph
    sig_para.clear()

    # Add the signature image
    run = sig_para.add_run()
    run.add_picture(signature_img_path, width=Inches(1.5))

    # Add the date text in the same paragraph
    sig_para.add_run("  Date: ______________")

    # Save the document
    doc.save(file_path)
    print(f"  ✓ Updated and saved successfully")

    return True

def main():
    """
    Main function to process all files.
    """
    print("Starting signature update process...")
    print(f"Signature image: {signature_image}")
    print(f"Base directory: {base_dir}")
    print(f"Total files to process: {len(files)}")

    # Check if signature image exists
    if not os.path.exists(signature_image):
        print(f"\nERROR: Signature image not found at {signature_image}")
        return

    success_count = 0
    failed_files = []

    for filename in files:
        file_path = os.path.join(base_dir, filename)

        if not os.path.exists(file_path):
            print(f"\nWARNING: File not found: {filename}")
            failed_files.append(filename)
            continue

        try:
            if update_signature(file_path, signature_image):
                success_count += 1
            else:
                failed_files.append(filename)
        except Exception as e:
            print(f"  ERROR: {str(e)}")
            failed_files.append(filename)

    # Summary
    print("\n" + "="*60)
    print("SUMMARY")
    print("="*60)
    print(f"Total files processed: {len(files)}")
    print(f"Successfully updated: {success_count}")
    print(f"Failed: {len(failed_files)}")

    if failed_files:
        print("\nFailed files:")
        for f in failed_files:
            print(f"  - {f}")
    else:
        print("\n✓ All 13 files updated successfully!")

if __name__ == "__main__":
    main()
