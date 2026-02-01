#!/usr/bin/env python3
"""
Script to debug the signature sections in Entertainment Agreement documents.
"""

from docx import Document

def debug_contract(file_path, name_search):
    """Debug the contract signature section."""
    print(f"\n=== Debugging: {file_path} ===")

    doc = Document(file_path)

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if name_search in text:
            print(f"\nParagraph {i}: '{text}'")
            # Show the next 5 paragraphs
            for j in range(i+1, min(i+6, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip()
                print(f"Paragraph {j}: '{next_text}'")
            break

def main():
    """Main function to debug both contracts."""

    boukou_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Boukou_Groove_2026.docx"
    killer_robot_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Killer_Robot_Army_2026_1.docx"

    debug_contract(boukou_path, "Name: Jennifer Sundal")
    debug_contract(killer_robot_path, "Name: Scott Rockwood")

if __name__ == "__main__":
    main()
