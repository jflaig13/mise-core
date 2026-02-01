#!/usr/bin/env python3
"""
Script to update Entertainment Agreement documents with performer names.
"""

import os
from docx import Document
import re

# Base directory containing the contracts
BASE_DIR = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/"

# List of contract files to process
CONTRACT_FILES = [
    "Entertainment_Agreement_Alex_Napier_Band_2026.docx",
    "Entertainment_Agreement_Black_Creek_String_Band_2026_1.docx",
    "Entertainment_Agreement_Blues_Old_Stand_Trio_2026.docx",
    "Entertainment_Agreement_Boukou_Groove_2026.docx",
    "Entertainment_Agreement_Dawghaus_Jazz_Band_2026.docx",
    "Entertainment_Agreement_Eric_Stockton_Band_2026.docx",
    "Entertainment_Agreement_Goodbye_To_Sunshine_2026.docx",
    "Entertainment_Agreement_John_Hart_Project_2026.docx",
    "Entertainment_Agreement_Killer_Robot_Army_2026_1.docx",
    "Entertainment_Agreement_Lips_Manly_2026.docx",
    "Entertainment_Agreement_Shenaniganza_2026.docx",
    "Entertainment_Agreement_Stevie_Monce_2026.docx",
    "Entertainment_Agreement_The_Unnamed_Band_2026.docx",
]

def extract_performer_name(doc):
    """
    Extract the performer name from the Parties Involved section.
    Looking for pattern: and [Name] ("Performer")
    """
    for para in doc.paragraphs:
        text = para.text
        # Look for pattern: and [something] ("Performer")
        # This captures the name between "and" and ("Performer")
        match = re.search(r'\band\s+([^(]+?)\s+\("Performer"\)', text)
        if match:
            return match.group(1).strip()
    return None

def update_performer_name_field(doc, performer_name):
    """
    Find and update the Performer's Name field in the signature section.
    Looking for "Name: ______________" or "Name: [existing text]"
    in the Artist/Performer Signature section.
    """
    updated = False
    performer_section_found = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Identify when we're in the Performer/Artist signature section
        if "Artist/Performer Signature" in text or "Performer Signature" in text:
            performer_section_found = True
            continue

        # If we find Organizer Signature, we've passed the performer section
        if "Organizer Signature" in text:
            performer_section_found = False
            continue

        # Look for the Name field (with underscores or existing text)
        # Skip if it contains "Jon Flaig" (that's the organizer field)
        if performer_section_found and text.startswith("Name:") and "Jon Flaig" not in text:
            # Replace the entire text with the performer name
            para.text = f"Name: {performer_name}"
            updated = True
            break

    return updated

def process_contract(file_path):
    """
    Process a single contract file.
    """
    try:
        # Open the document
        doc = Document(file_path)

        # Extract performer name
        performer_name = extract_performer_name(doc)
        if not performer_name:
            return None, "Could not extract performer name"

        # Update the performer name field
        updated = update_performer_name_field(doc, performer_name)
        if not updated:
            return performer_name, "Could not find/update Name field"

        # Save the document
        doc.save(file_path)

        return performer_name, "Success"

    except Exception as e:
        return None, f"Error: {str(e)}"

def main():
    """
    Main function to process all contracts.
    """
    print("Processing Entertainment Agreement Contracts")
    print("=" * 60)

    results = []

    for i, filename in enumerate(CONTRACT_FILES, 1):
        file_path = os.path.join(BASE_DIR, filename)
        print(f"\n{i}. Processing: {filename}")

        if not os.path.exists(file_path):
            print(f"   File not found!")
            results.append((filename, None, "File not found"))
            continue

        performer_name, status = process_contract(file_path)
        results.append((filename, performer_name, status))

        if status == "Success":
            print(f"   Updated with performer name: {performer_name}")
        else:
            print(f"   {status}")
            if performer_name:
                print(f"   Found performer: {performer_name}")

    # Summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)

    for filename, performer_name, status in results:
        if status == "Success":
            print(f"Contract: {filename}")
            print(f"  Performer: {performer_name}")
            print()
        else:
            print(f"Failed: {filename}")
            print(f"  Status: {status}")
            print()

    successful = sum(1 for _, _, status in results if status == "Success")
    print(f"\nTotal: {successful}/{len(CONTRACT_FILES)} contracts updated successfully")

if __name__ == "__main__":
    main()
