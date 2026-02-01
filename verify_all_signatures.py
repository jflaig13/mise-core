#!/usr/bin/env python3
"""
Script to verify all 13 documents have the signature properly embedded.
"""

from docx import Document
import os

base_dir = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026"

files = [
    "Entertainment_Agreement_Alex_Napier_Band_2026.docx",
    "Entertainment_Agreement_Black_Creek_String_Band_2026_1.docx",
    "Entertainment_Agreement_Blues_Old_Stand_Trio_2026.docx",
    "Entertainment_Agreement_Boukou_Groove_2026.docx",
    "Entertainment_Agreement_Dawghaus_Jazz_Band_2026.docx",
    "Entertainment_Agreement_Eric_Stockton_Band_2026.docx",
    "Entertainment_Agreement_Goodbye_To_Sunshine_2026.docx",
    "Entertainment_Agreement_John_Hart_Project_2026.docx",
    "Entertainment_Agreement_Killer_Robot_Army_2026_1.docx",
    "Entertainment_Agreement_Lips_Manly_2026.docx",
    "Entertainment_Agreement_Shenaniganza_2026.docx",
    "Entertainment_Agreement_Stevie_Monce_2026.docx",
    "Entertainment_Agreement_The_Unnamed_Band_2026.docx",
]

def verify_signature(file_path):
    """Verify that a document has the signature image properly embedded."""
    doc = Document(file_path)

    # First, find "Name: Jon Flaig"
    jon_flaig_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "Name: Jon Flaig" in para.text:
            jon_flaig_idx = i
            break

    if jon_flaig_idx is None:
        return False, False, "Name: Jon Flaig not found"

    # Find the organizer line
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if "On behalf of Bloom's Restaurant, LLC d/b/a Papa Surf Burger Bar:" in text:
            # The signature paragraph is the one right before "Name: Jon Flaig"
            if i < jon_flaig_idx:
                sig_para = doc.paragraphs[jon_flaig_idx - 1]

                # Check for image in runs
                has_image = False
                for run in sig_para.runs:
                    if run._element.xpath('.//pic:pic'):
                        has_image = True
                        break

                # Check for date text
                has_date_text = "Date:" in sig_para.text

                return has_image, has_date_text, sig_para.text

    return False, False, "Signature paragraph not found"

print("Verifying all 13 entertainment agreements...\n")
print("="*80)

all_success = True
for filename in files:
    file_path = os.path.join(base_dir, filename)
    has_image, has_date, para_text = verify_signature(file_path)

    status = "✓" if (has_image and has_date) else "✗"
    all_success = all_success and (has_image and has_date)

    print(f"{status} {filename}")
    print(f"   Image: {'Yes' if has_image else 'NO'} | Date text: {'Yes' if has_date else 'NO'}")
    if not (has_image and has_date):
        print(f"   Para text: '{para_text}'")

print("="*80)
if all_success:
    print("\n✓ SUCCESS: All 13 files verified with signature image and date text!")
else:
    print("\n✗ FAILED: Some files are missing signature image or date text")
