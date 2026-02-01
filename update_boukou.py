#!/usr/bin/env python3
"""
Script to update Boukou Groove contract signature section.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

def update_boukou_groove_contract(file_path):
    """Update the Boukou Groove contract signature section."""
    print(f"\n=== Processing: {file_path} ===")

    doc = Document(file_path)
    found = False

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if "Name: Jennifer Sundal" in text:
            # Check if there's already a Title paragraph after it
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if "Title: Manager, Sound Production Inc." in next_text:
                    print(f"Already updated! Found at paragraph {i}")
                    return True

            print(f"Found target text at paragraph {i}: '{text}'")

            # Store the formatting from the original paragraph
            original_runs = paragraph.runs
            font_name = None
            font_size = None
            font_bold = None
            font_color = None

            if original_runs:
                font_name = original_runs[0].font.name
                font_size = original_runs[0].font.size
                font_bold = original_runs[0].font.bold
                font_color = original_runs[0].font.color.rgb if original_runs[0].font.color.rgb else None

            print(f"Original formatting - Font: {font_name}, Size: {font_size}, Bold: {font_bold}")

            # Create a new paragraph element and insert it after the current one
            parent = paragraph._element.getparent()
            new_p = OxmlElement('w:p')
            parent.insert(parent.index(paragraph._element) + 1, new_p)
            
            # Create the new paragraph object
            new_paragraph = Paragraph(new_p, paragraph._parent)
            
            # Add the text with formatting
            new_run = new_paragraph.add_run("Title: Manager, Sound Production Inc.")
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
            if font_bold is not None:
                new_run.font.bold = font_bold
            if font_color:
                new_run.font.color.rgb = font_color

            print(f"Inserted new paragraph: 'Title: Manager, Sound Production Inc.'")

            found = True
            break

    if found:
        doc.save(file_path)
        print(f"✓ Successfully saved changes to {file_path}")
        return True
    else:
        print(f"✗ Could not find target text in {file_path}")
        return False

if __name__ == "__main__":
    boukou_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Boukou_Groove_2026.docx"
    
    try:
        result = update_boukou_groove_contract(boukou_path)
        print("\n" + "="*60)
        if result:
            print("✓ SUCCESS")
        else:
            print("✗ FAILED")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
