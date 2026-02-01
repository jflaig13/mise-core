#!/usr/bin/env python3
"""
Script to update signature sections in Entertainment Agreement documents.
"""

from docx import Document
from docx.oxml import OxmlElement
import sys

def update_boukou_groove_contract(file_path):
    """Update the Boukou Groove contract signature section."""
    print(f"\n=== Processing: {file_path} ===")

    doc = Document(file_path)
    found = False

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if "Name: Jennifer Sundal" in text and i > 0:
            # Check if there's already a Title paragraph after it
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if "Title: Manager, Sound Production Inc." in next_text:
                    print(f"Found target text already updated at paragraph {i}")
                    found = True
                    break

            print(f"Found target text at paragraph {i}: '{text}'")

            # Store the formatting from the original paragraph
            original_runs = paragraph.runs
            font_name = None
            font_size = None
            font_bold = None
            font_color = None

            if original_runs:
                font_name = original_runs[0].font.name
                font_size = original_runs[0].font.size
                font_bold = original_runs[0].font.bold
                font_color = original_runs[0].font.color.rgb if original_runs[0].font.color.rgb else None

            print(f"Original formatting - Font: {font_name}, Size: {font_size}, Bold: {font_bold}")

            # Create a new paragraph element and insert it after the current one
            new_para = paragraph._element.addnext(paragraph._element.makeelement('w:p', nsmap=paragraph._element.nsmap))
            new_paragraph = Document()._element.body.append(new_para)
            
            # Simpler approach: get the parent and insert after
            parent = paragraph._element.getparent()
            new_p = OxmlElement('w:p')
            parent.insert(parent.index(paragraph._element) + 1, new_p)
            
            # Create the new paragraph object
            from docx.text.paragraph import Paragraph
            new_paragraph = Paragraph(new_p, paragraph._parent)
            
            # Add the text with formatting
            new_run = new_paragraph.add_run("Title: Manager, Sound Production Inc.")
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
            if font_bold is not None:
                new_run.font.bold = font_bold
            if font_color:
                new_run.font.color.rgb = font_color

            print(f"Inserted new paragraph: 'Title: Manager, Sound Production Inc.'")

            found = True
            break

        # Also check for the old format
        if "Name: Sound Production Inc., By: Jennifer Sundal, Manager" in text:
            print(f"Found old format at paragraph {i}: '{text}'")

            # Store the formatting from the original paragraph
            original_runs = paragraph.runs
            font_name = None
            font_size = None
            font_bold = None
            font_color = None

            if original_runs:
                font_name = original_runs[0].font.name
                font_size = original_runs[0].font.size
                font_bold = original_runs[0].font.bold
                font_color = original_runs[0].font.color.rgb if original_runs[0].font.color.rgb else None

            print(f"Original formatting - Font: {font_name}, Size: {font_size}, Bold: {font_bold}")

            # Clear the paragraph and set new text
            paragraph.clear()
            run = paragraph.add_run("Name: Jennifer Sundal")

            # Apply the same formatting
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if font_bold is not None:
                run.font.bold = font_bold
            if font_color:
                run.font.color.rgb = font_color

            print(f"Replaced with: 'Name: Jennifer Sundal'")

            # Create a new paragraph element and insert it after the current one
            parent = paragraph._element.getparent()
            new_p = OxmlElement('w:p')
            parent.insert(parent.index(paragraph._element) + 1, new_p)
            
            # Create the new paragraph object
            from docx.text.paragraph import Paragraph
            new_paragraph = Paragraph(new_p, paragraph._parent)
            
            # Add the text with formatting
            new_run = new_paragraph.add_run("Title: Manager, Sound Production Inc.")
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
            if font_bold is not None:
                new_run.font.bold = font_bold
            if font_color:
                new_run.font.color.rgb = font_color

            print(f"Inserted new paragraph: 'Title: Manager, Sound Production Inc.'")

            found = True
            break

    if found:
        doc.save(file_path)
        print(f"✓ Successfully saved changes to {file_path}")
        return True
    else:
        print(f"✗ Could not find target text in {file_path}")
        return False

def update_killer_robot_army_contract(file_path):
    """Update the Killer Robot Army contract signature section."""
    print(f"\n=== Processing: {file_path} ===")

    doc = Document(file_path)
    found = False

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if "Name: Scott Rockwood" in text and "LLC" not in text:
            # Check if there's already a Title paragraph after it
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if "Title: Member, Scott Rockwood, LLC" in next_text:
                    print(f"Found target text already updated at paragraph {i}")
                    found = True
                    break

            print(f"Found target text at paragraph {i}: '{text}'")

            # Store the formatting from the original paragraph
            original_runs = paragraph.runs
            font_name = None
            font_size = None
            font_bold = None
            font_color = None

            if original_runs:
                font_name = original_runs[0].font.name
                font_size = original_runs[0].font.size
                font_bold = original_runs[0].font.bold
                font_color = original_runs[0].font.color.rgb if original_runs[0].font.color.rgb else None

            print(f"Original formatting - Font: {font_name}, Size: {font_size}, Bold: {font_bold}")

            # Create a new paragraph element and insert it after the current one
            parent = paragraph._element.getparent()
            new_p = OxmlElement('w:p')
            parent.insert(parent.index(paragraph._element) + 1, new_p)
            
            # Create the new paragraph object
            from docx.text.paragraph import Paragraph
            new_paragraph = Paragraph(new_p, paragraph._parent)
            
            # Add the text with formatting
            new_run = new_paragraph.add_run("Title: Member, Scott Rockwood, LLC")
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
            if font_bold is not None:
                new_run.font.bold = font_bold
            if font_color:
                new_run.font.color.rgb = font_color

            print(f"Inserted new paragraph: 'Title: Member, Scott Rockwood, LLC'")

            found = True
            break

        # Also check for the old format
        if "Name: Scott Rockwood, LLC, By: Scott Rockwood, Member" in text:
            print(f"Found old format at paragraph {i}: '{text}'")

            # Store the formatting from the original paragraph
            original_runs = paragraph.runs
            font_name = None
            font_size = None
            font_bold = None
            font_color = None

            if original_runs:
                font_name = original_runs[0].font.name
                font_size = original_runs[0].font.size
                font_bold = original_runs[0].font.bold
                font_color = original_runs[0].font.color.rgb if original_runs[0].font.color.rgb else None

            print(f"Original formatting - Font: {font_name}, Size: {font_size}, Bold: {font_bold}")

            # Clear the paragraph and set new text
            paragraph.clear()
            run = paragraph.add_run("Name: Scott Rockwood")

            # Apply the same formatting
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if font_bold is not None:
                run.font.bold = font_bold
            if font_color:
                run.font.color.rgb = font_color

            print(f"Replaced with: 'Name: Scott Rockwood'")

            # Create a new paragraph element and insert it after the current one
            parent = paragraph._element.getparent()
            new_p = OxmlElement('w:p')
            parent.insert(parent.index(paragraph._element) + 1, new_p)
            
            # Create the new paragraph object
            from docx.text.paragraph import Paragraph
            new_paragraph = Paragraph(new_p, paragraph._parent)
            
            # Add the text with formatting
            new_run = new_paragraph.add_run("Title: Member, Scott Rockwood, LLC")
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
            if font_bold is not None:
                new_run.font.bold = font_bold
            if font_color:
                new_run.font.color.rgb = font_color

            print(f"Inserted new paragraph: 'Title: Member, Scott Rockwood, LLC'")

            found = True
            break

    if found:
        doc.save(file_path)
        print(f"✓ Successfully saved changes to {file_path}")
        return True
    else:
        print(f"✗ Could not find target text in {file_path}")
        return False

def main():
    """Main function to update both contracts."""

    boukou_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Boukou_Groove_2026.docx"
    killer_robot_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Killer_Robot_Army_2026_1.docx"

    results = []

    # Update Boukou Groove contract
    try:
        result1 = update_boukou_groove_contract(boukou_path)
        results.append(("Boukou Groove", result1))
    except Exception as e:
        print(f"Error processing Boukou Groove contract: {e}")
        import traceback
        traceback.print_exc()
        results.append(("Boukou Groove", False))

    # Update Killer Robot Army contract
    try:
        result2 = update_killer_robot_army_contract(killer_robot_path)
        results.append(("Killer Robot Army", result2))
    except Exception as e:
        print(f"Error processing Killer Robot Army contract: {e}")
        import traceback
        traceback.print_exc()
        results.append(("Killer Robot Army", False))

    # Summary
    print("\n" + "="*60)
    print("SUMMARY")
    print("="*60)
    for contract_name, success in results:
        status = "✓ SUCCESS" if success else "✗ FAILED"
        print(f"{contract_name}: {status}")

    all_success = all(result[1] for result in results)
    sys.exit(0 if all_success else 1)

if __name__ == "__main__":
    main()
