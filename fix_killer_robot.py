#!/usr/bin/env python3
"""
Script to fix the Killer Robot Army document.
"""

from docx import Document
from docx.shared import Inches

file_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Killer_Robot_Army_2026_1.docx"
signature_image = "/Users/jonathanflaig/Signature/JonFlaig_Signature.png"

print(f"Fixing: Entertainment_Agreement_Killer_Robot_Army_2026_1.docx")

# Load the document
doc = Document(file_path)

# Find "Name: Jon Flaig" paragraph
jon_flaig_idx = None
for i, para in enumerate(doc.paragraphs):
    if "Name: Jon Flaig" in para.text:
        jon_flaig_idx = i
        print(f"Found 'Name: Jon Flaig' at paragraph {i}")
        break

if jon_flaig_idx is None:
    print("ERROR: Could not find 'Name: Jon Flaig'")
    exit(1)

# Find the organizer line and the signature paragraph
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    if "On behalf of Bloom's Restaurant, LLC d/b/a Papa Surf Burger Bar:" in text:
        print(f"Found organizer line at paragraph {i}")

        # The signature paragraph should be right before "Name: Jon Flaig"
        sig_para_idx = jon_flaig_idx - 1
        sig_para = doc.paragraphs[sig_para_idx]

        print(f"Signature paragraph is at index {sig_para_idx}")
        print(f"Current content: '{sig_para.text}'")

        # Clear the paragraph
        sig_para.clear()

        # Add the signature image
        run = sig_para.add_run()
        run.add_picture(signature_image, width=Inches(1.5))

        # Add the date text in the same paragraph
        sig_para.add_run("  Date: ______________")

        # Save the document
        doc.save(file_path)
        print("✓ Updated and saved successfully")
        break
