#!/usr/bin/env python3
"""
Debug script to examine the Killer Robot Army document structure.
"""

from docx import Document

file_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Killer_Robot_Army_2026_1.docx"

doc = Document(file_path)

print("Looking for organizer signature section...\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # Look for the organizer section
    if "Bloom's Restaurant" in text or "Papa Surf" in text or "ORGANIZER" in text:
        print(f"Paragraph {i}: '{text}'")
        # Show surrounding paragraphs
        for j in range(-2, 5):
            idx = i + j
            if 0 <= idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]
                marker = " <<< THIS ONE" if j == 0 else ""
                print(f"  [{idx}] '{p.text}'{marker}")
        print()

print("\nSearching for 'Name: Jon Flaig'...")
for i, para in enumerate(doc.paragraphs):
    if "Jon Flaig" in para.text:
        print(f"Found at paragraph {i}: '{para.text}'")
        # Show surrounding paragraphs
        for j in range(-3, 2):
            idx = i + j
            if 0 <= idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]
                marker = " <<< THIS ONE" if j == 0 else ""
                print(f"  [{idx}] '{p.text}'{marker}")
