#!/usr/bin/env python3
"""
Script to verify the signature was properly added to a sample document.
"""

from docx import Document
import os

# Check one sample file
sample_file = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Alex_Napier_Band_2026.docx"

print(f"Verifying: {os.path.basename(sample_file)}\n")

doc = Document(sample_file)

# Find the organizer signature section
found_organizer = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    if "On behalf of Bloom's Restaurant, LLC d/b/a Papa Surf Burger Bar:" in text:
        print(f"Paragraph {i}: {text}")
        found_organizer = True
        continue

    if found_organizer and i < len(doc.paragraphs) - 1:
        # Check next 3 paragraphs
        for j in range(3):
            if i + j < len(doc.paragraphs):
                p = doc.paragraphs[i + j]
                print(f"Paragraph {i + j}: Text='{p.text}'")

                # Check for images in runs
                has_image = False
                for run in p.runs:
                    if run._element.xpath('.//pic:pic'):
                        has_image = True
                        print(f"  -> Contains embedded image")

                if not has_image and len(p.runs) > 0:
                    print(f"  -> No image found, {len(p.runs)} runs")
        break

print("\n✓ Verification complete")
