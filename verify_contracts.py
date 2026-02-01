#!/usr/bin/env python3
"""
Script to verify the signature sections in Entertainment Agreement documents.
"""

from docx import Document

def verify_contract(file_path, expected_name, expected_title):
    """Verify the contract signature section."""
    print(f"\n=== Verifying: {file_path} ===")

    doc = Document(file_path)
    found_name = False
    found_title = False

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if expected_name in text:
            print(f"✓ Found at paragraph {i}: '{text}'")
            found_name = True
            # Check if the next paragraph has the title
            if i + 1 < len(doc.paragraphs):
                next_text = doc.paragraphs[i + 1].text.strip()
                if expected_title in next_text:
                    print(f"✓ Found at paragraph {i + 1}: '{next_text}'")
                    found_title = True
            break

    if found_name and found_title:
        print("✓ Verification PASSED")
        return True
    else:
        print("✗ Verification FAILED")
        return False

def main():
    """Main function to verify both contracts."""

    boukou_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Boukou_Groove_2026.docx"
    killer_robot_path = "/Users/jonathanflaig/Library/CloudStorage/OneDrive-Personal/Papa Surf/Live Music/Contracts/2026/Entertainment_Agreement_Killer_Robot_Army_2026_1.docx"

    results = []

    # Verify Boukou Groove contract
    try:
        result1 = verify_contract(boukou_path, "Name: Jennifer Sundal", "Title: Manager, Sound Production Inc.")
        results.append(("Boukou Groove", result1))
    except Exception as e:
        print(f"Error verifying Boukou Groove contract: {e}")
        results.append(("Boukou Groove", False))

    # Verify Killer Robot Army contract
    try:
        result2 = verify_contract(killer_robot_path, "Name: Scott Rockwood", "Title: Member, Scott Rockwood, LLC")
        results.append(("Killer Robot Army", result2))
    except Exception as e:
        print(f"Error verifying Killer Robot Army contract: {e}")
        results.append(("Killer Robot Army", False))

    # Summary
    print("\n" + "="*60)
    print("VERIFICATION SUMMARY")
    print("="*60)
    for contract_name, success in results:
        status = "✓ PASS" if success else "✗ FAIL"
        print(f"{contract_name}: {status}")

if __name__ == "__main__":
    main()
